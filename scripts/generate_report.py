"""Build the formatted technical report (.docx) from the authored content.

Applies every formatting rule mandated by the course brief:

* Times New Roman 12pt body, bold Heading styles (so they populate the ToC).
* 1.5 line spacing, 1 inch margins on all sides.
* Cover page, abstract, an auto-updating Word Table of Contents, plus manual
  lists of figures and tables.
* Figures numbered and captioned below; tables numbered and captioned above;
  equations numbered and right-aligned; code in Courier New 10pt on a grey
  background with line numbers.
* Page numbers bottom-centre, restarting at 1 from the Introduction.

Usage (after seeding + figures)::

    python -m scripts.seed
    python -m scripts.figures
    python -m scripts.generate_report      # -> docs/Technical_Report.docx
"""

from __future__ import annotations

import glob
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from scripts import report_content as C

OUTPUT = os.path.join("docs", "Technical_Report.docx")
FIG_DIR = os.path.join("docs", "figures")

# Running counters for numbered captions/equations.
_counters = {"figure": 0, "table": 0, "equation": 0}
_figure_index: list[str] = []
_table_index: list[str] = []


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------
def _add_field(paragraph, instruction: str, placeholder: str = ""):
    """Insert a Word field (e.g. PAGE or TOC) into a paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, text, end):
        run._r.append(el)


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _restart_page_numbers(section) -> None:
    """Make the given section restart page numbering at 1."""
    sect_pr = section._sectPr
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "1")
    sect_pr.append(pg)


# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------
def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    # Ensure the font applies to all script ranges.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1 (section titles) - 14pt bold; Heading 2 (sub-headings) - 12pt bold.
    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------
def _para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        doc.add_paragraph(item, style=style)


def _equation(doc, text):
    _counters["equation"] += 1
    n = _counters["equation"]
    p = doc.add_paragraph()
    # Tab stop at the right margin so the equation number is flush right.
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(text)
    run.italic = True
    p.add_run(f"\t({n})")


def _figure(doc, filename, caption):
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _counters["figure"] += 1
    n = _counters["figure"]
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {n}: {caption}")
    run.italic = True
    run.font.size = Pt(10)
    _figure_index.append(f"Figure {n}: {caption}")


def _table(doc, caption, headers, rows):
    _counters["table"] += 1
    n = _counters["table"]
    cap = doc.add_paragraph()
    run = cap.add_run(f"Table {n}: {caption}")
    run.bold = True
    run.font.size = Pt(10)
    _table_index.append(f"Table {n}: {caption}")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    doc.add_paragraph()


def _code(doc, caption, code_text):
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, "EEEEEE")
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = 1.0
    lines = code_text.split("\n")
    width = len(str(len(lines)))
    for idx, line in enumerate(lines, start=1):
        if idx > 1:
            para.add_run().add_break()
        run = para.add_run(f"{str(idx).rjust(width)} | {line}")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        # Courier New for all script ranges.
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), "Courier New")
    doc.add_paragraph()


def _latest_evidence() -> str:
    files = sorted(glob.glob(os.path.join("test_evidence", "report_*.txt")))
    if not files:
        return "(No captured test evidence found. Run python -m scripts.capture_test_results.)"
    with open(files[-1], "r", encoding="utf-8", errors="replace") as fh:
        content = fh.read()
    # Trim to a reasonable length for the appendix.
    return content[:6000]


def _render_blocks(doc, blocks):
    for block in blocks:
        kind = block[0]
        if kind == "para":
            _para(doc, block[1])
        elif kind == "sub":
            doc.add_heading(block[1], level=2)
        elif kind == "bullet":
            _bullets(doc, block[1], numbered=False)
        elif kind == "numbered":
            _bullets(doc, block[1], numbered=True)
        elif kind == "figure":
            _figure(doc, block[1], block[2])
        elif kind == "table":
            _table(doc, block[1], block[2], block[3])
        elif kind == "equation":
            _equation(doc, block[1])
        elif kind == "code":
            _code(doc, block[1], block[2])
        elif kind == "evidence":
            _code(doc, "Captured pytest output (excerpt)", _latest_evidence())


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------
def _cover_page(doc):
    m = C.META

    def centered(text, size=12, bold=False, space_before=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if space_before:
            p.paragraph_format.space_before = Pt(space_before)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    centered(m["logo_placeholder"], size=12, space_before=24)
    centered(m["university"], size=16, bold=True, space_before=24)
    centered(m["department"], size=13)
    centered(m["title"], size=18, bold=True, space_before=48)
    centered(f"{m['course_code']} - {m['course_name']}", size=13, space_before=36)
    centered("Submitted by:", size=12, bold=True, space_before=36)
    for member in m["team"]:
        centered(member, size=12)
    centered(f"Supervisor: {m['supervisor']}", size=12, bold=True, space_before=24)
    centered(f"Date: {m['date']}", size=12)


def _abstract(doc):
    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Abstract")
    run.bold = True
    run.font.size = Pt(14)
    p = _para(doc, C.ABSTRACT)
    word_count = len(C.ABSTRACT.split())
    note = doc.add_paragraph()
    note_run = note.add_run(f"(Word count: {word_count})")
    note_run.italic = True
    note_run.font.size = Pt(9)


def _table_of_contents(doc):
    doc.add_page_break()
    heading = doc.add_paragraph()
    run = heading.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    instr = doc.add_paragraph()
    irun = instr.add_run(
        "(In Microsoft Word, right-click the table below and choose "
        "\"Update Field\" to populate page numbers.)"
    )
    irun.italic = True
    irun.font.size = Pt(9)
    toc_para = doc.add_paragraph()
    _add_field(toc_para, 'TOC \\o "1-2" \\h \\z \\u', "Update this field to build the Table of Contents.")


def _lists_of_figures_tables(doc):
    """Manual lists of figures and tables (populated from build-time counters)."""
    doc.add_page_break()
    h1 = doc.add_paragraph()
    r1 = h1.add_run("List of Figures")
    r1.bold = True
    r1.font.size = Pt(14)
    if _figure_index:
        for entry in _figure_index:
            doc.add_paragraph(entry)
    else:
        doc.add_paragraph("(none)")

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    r2 = h2.add_run("List of Tables")
    r2.bold = True
    r2.font.size = Pt(14)
    if _table_index:
        for entry in _table_index:
            doc.add_paragraph(entry)
    else:
        doc.add_paragraph("(none)")


def _start_body_section(doc):
    """Begin a new section for the body so page numbers can restart at 1."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    _restart_page_numbers(section)
    # Footer: centred page number, not linked to the front-matter section.
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(fp, "PAGE", "1")
    return section


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------
def build() -> str:
    doc = Document()
    _configure_styles(doc)

    # We render the body first (into a deferred list) so that the lists of
    # figures/tables are complete before we emit them in the front matter.
    # Simpler approach: render front matter placeholders, body, then because
    # python-docx is append-only, we build body counters during the body pass.
    # To keep lists correct, we therefore emit the front matter, then a body
    # section, and put the lists of figures/tables AFTER the ToC but we need
    # counts first. We resolve this by doing a two-pass: pre-count by walking
    # the content, then build linearly.

    # Pre-count figures/tables for the lists.
    for section in C.SECTIONS:
        for block in section["blocks"]:
            if block[0] == "figure":
                _counters["figure"] += 1
                _figure_index.append(f"Figure {_counters['figure']}: {block[2]}")
            elif block[0] == "table":
                _counters["table"] += 1
                _table_index.append(f"Table {_counters['table']}: {block[2]}")
    # Reset counters; the real render will renumber identically.
    _counters["figure"] = 0
    _counters["table"] = 0

    # Front matter (no page numbers).
    _cover_page(doc)
    _abstract(doc)
    _table_of_contents(doc)
    _lists_of_figures_tables(doc)

    # Body section (page numbers restart at 1).
    _start_body_section(doc)

    for section in C.SECTIONS:
        doc.add_heading(f"{section['number']}. {section['title']}", level=1)
        _render_blocks(doc, section["blocks"])

    # References.
    doc.add_heading("References", level=1)
    for i, ref in enumerate(C.REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.add_run(f"[{i}] {ref}")

    # Appendices.
    for appendix in C.APPENDICES:
        doc.add_heading(appendix["title"], level=1)
        _render_blocks(doc, appendix["blocks"])

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    out_path = OUTPUT
    try:
        doc.save(out_path)
    except PermissionError:
        # The canonical file is open in a viewer (Word/preview). Fall back to a
        # timestamped name so the build still succeeds; the user can rename it.
        from datetime import datetime

        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = OUTPUT.replace(".docx", f"_{stamp}.docx")
        doc.save(out_path)
        print(f"NOTE: {OUTPUT} was locked (open in a viewer?); wrote {out_path} instead.")
    print(f"Report written to {out_path}")
    print(f"  figures embedded: {_counters['figure']}, tables: {_counters['table']}, "
          f"equations: {_counters['equation']}")
    return out_path


if __name__ == "__main__":
    build()
